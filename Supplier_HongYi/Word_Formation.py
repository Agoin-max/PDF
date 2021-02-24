import sys

import docx
from win32com import client as wc
from openpyxl import Workbook


class TranDoc:
    def doc_docx(self, path):
        w = wc.Dispatch('Word.Application')
        doc = w.Documents.Open(path)
        docFile = path + "x"
        doc.SaveAs(docFile, 16)
        doc.Close()
        w.Quit()
        return path + "x"


class DocxExcel:
    def docx_excel(self, document):
        count = 0
        tables = []
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'C/NO'
        ws['B1'] = 'Brand'
        ws['C1'] = 'Description_c'
        ws['D1'] = 'Description_d'
        ws['E1'] = 'Customer PO'
        ws['F1'] = 'QTY'
        ws['G1'] = 'NW'
        ws['H1'] = 'GW'

        for item in document.paragraphs:
            print(item.text)

        sys.exit()
        total = len(document.tables)
        # print("总共", total, "个表格等待处理，请喝杯咖啡等待许久...")
        for index in range(0, total):
            table = []
            for row in document.tables[index].rows:
                line = []
                for grid in row.cells:
                    line.append(grid.text)
                table.append(line)
                ws.append(line)
            count = count + 1
            # print("第", count, "个表格正在处理...剩余", total - count + 1, "个表格", "\n")
            tables.append(table)
        wb.save("../ExcelFile/HongYi.xlsx")
        # print("表格处理完成...")


path = r"C:\Users\windo\Desktop\【类型3】13家供应商\智微\13弘憶\IN2-20B0285 PK.docx"
# docFile = TranDoc().doc_docx(path)
document = docx.Document(path)
DocxExcel().docx_excel(document)
